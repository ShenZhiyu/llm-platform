"""智能写作 Word 模板处理服务。

负责模板上传、.doc 转 .docx、解析 <title>/<body>、渲染生成 Word 文件。
"""

import hashlib
import json
import re
import shutil
import subprocess
from dataclasses import dataclass
from pathlib import Path
from zipfile import BadZipFile

from docx import Document
from docx.shared import Pt
from fastapi import UploadFile

from app.api.utils import new_id


WRITING_STORAGE = Path("storage") / "writing"
TEMPLATE_STORAGE = WRITING_STORAGE / "templates"
DOCUMENT_STORAGE = WRITING_STORAGE / "documents"

TITLE_BLOCK_PATTERN = re.compile(r"<title\b[^>]*>(?P<content>.*?)</title>", re.IGNORECASE | re.DOTALL)
BODY_BLOCK_PATTERN = re.compile(r"<body\b[^>]*>(?P<content>.*?)</body>", re.IGNORECASE | re.DOTALL)
PLACEHOLDER_PATTERN = re.compile(r"\{\{\s*(title|body)\s*\}\}", re.IGNORECASE)


class WritingServiceError(Exception):
    """模板解析、转换或渲染失败。"""

    pass


@dataclass
class StoredUpload:
    file_name: str
    file_path: str
    file_size: int
    content_hash: str


def ensure_writing_storage() -> None:
    """确保模板文件和生成文档的本地存储目录存在。"""
    TEMPLATE_STORAGE.mkdir(parents=True, exist_ok=True)
    DOCUMENT_STORAGE.mkdir(parents=True, exist_ok=True)


def file_sha256(path: Path) -> str:
    """计算文件 SHA-256，用于去重和版本记录。"""
    digest = hashlib.sha256()
    with path.open("rb") as file:
        for chunk in iter(lambda: file.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def convert_doc_to_docx(source: Path) -> Path:
    """.doc 模板通过 LibreOffice 转换为 .docx 后再解析。"""
    converter = shutil.which("soffice") or shutil.which("libreoffice")
    if converter is None:
        raise WritingServiceError("上传 .doc 模板需要服务器安装 LibreOffice，并确保 soffice/libreoffice 在 PATH 中。")
    result = subprocess.run(
        [converter, "--headless", "--convert-to", "docx", "--outdir", str(TEMPLATE_STORAGE), str(source)],
        capture_output=True,
        text=True,
        timeout=60,
        check=False,
    )
    converted = TEMPLATE_STORAGE / f"{source.stem}.docx"
    if result.returncode != 0 or not converted.exists():
        detail = (result.stderr or result.stdout or "").strip()
        raise WritingServiceError(f".doc 模板转换 .docx 失败{f'：{detail}' if detail else ''}")
    return converted


def save_template_upload(file: UploadFile) -> StoredUpload:
    """保存上传的 Word 模板并校验文件可解析。"""
    ensure_writing_storage()
    suffix = Path(file.filename or "").suffix.lower()
    if suffix not in {".doc", ".docx"}:
        raise WritingServiceError("当前仅支持上传 .doc 或 .docx Word 模板。")

    safe_name = Path(file.filename or f"template{suffix or '.docx'}").name
    target = TEMPLATE_STORAGE / f"{new_id('tplfile')}_{safe_name}"
    with target.open("wb") as output:
        shutil.copyfileobj(file.file, output)

    parse_target = target
    if suffix == ".doc":
        try:
            parse_target = convert_doc_to_docx(target)
        except WritingServiceError:
            target.unlink(missing_ok=True)
            raise

    try:
        Document(str(parse_target))
    except (BadZipFile, Exception) as exc:
        target.unlink(missing_ok=True)
        if parse_target != target:
            parse_target.unlink(missing_ok=True)
        raise WritingServiceError("Word 模板无法解析，请确认文件是有效的 .doc 或 .docx。") from exc

    return StoredUpload(
        file_name=safe_name,
        file_path=str(parse_target),
        file_size=target.stat().st_size,
        content_hash=file_sha256(parse_target),
    )


def _iter_paragraphs(document: Document):
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph
    for section in document.sections:
        for paragraph in section.header.paragraphs:
            yield paragraph
        for paragraph in section.footer.paragraphs:
            yield paragraph


def _iter_body_paragraphs(document: Document):
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def _paragraph_groups(document: Document):
    yield document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield cell.paragraphs


def clean_tag_content(value: str) -> str:
    lines = [line.strip() for line in value.replace("\r\n", "\n").replace("\r", "\n").split("\n")]
    while lines and not lines[0]:
        lines.pop(0)
    while lines and not lines[-1]:
        lines.pop()
    return "\n".join(lines)


def _text_from_document(document: Document) -> str:
    return "\n".join(paragraph.text for paragraph in _iter_paragraphs(document))


def _body_text_from_document(document: Document) -> str:
    return "\n".join(paragraph.text for paragraph in _iter_body_paragraphs(document))


def _field(key: str, label: str, placeholder: str, default_value: str, order: int) -> dict:
    return {
        "key": key,
        "label": label,
        "placeholder": placeholder,
        "type": key,
        "editable": True,
        "formatEditable": False,
        "defaultValue": default_value,
        "removable": False,
        "addable": False,
        "order": order,
    }


def extract_template_metadata(file_path: str) -> tuple[list[dict], str]:
    """Parse only <title> and <body>. Section tags are intentionally unsupported."""
    document = Document(file_path)
    text = _body_text_from_document(document)
    title_match = TITLE_BLOCK_PATTERN.search(text)
    body_match = BODY_BLOCK_PATTERN.search(text)

    fields = [
        _field("title", "标题", "<title>", clean_tag_content(title_match.group("content")) if title_match else "", 0),
        _field("body", "正文", "<body>", clean_tag_content(body_match.group("content")) if body_match else "", 1),
    ]
    if not title_match and "{{title}}" in text:
        fields[0]["placeholder"] = "{{title}}"
    if not body_match and "{{body}}" in text:
        fields[1]["placeholder"] = "{{body}}"

    preview = re.sub(r"</?(title|body)\b[^>]*>", "", text, flags=re.IGNORECASE)
    return fields, "\n".join(clean_tag_content(preview).splitlines()[:160])


def default_format_config() -> dict:
    return {
        "titleFont": "黑体",
        "bodyFont": "仿宋",
        "titleFontSize": "二号",
        "bodyFontSize": "小四",
        "fontSize": "小四",
        "lineSpacing": "1.5",
        "allowUserFormat": False,
    }


def json_dumps(value) -> str:
    return json.dumps(value, ensure_ascii=False)


def json_loads(value: str | None, fallback):
    if not value:
        return fallback
    try:
        parsed = json.loads(value)
    except json.JSONDecodeError:
        return fallback
    return parsed if parsed is not None else fallback


def content_from_fields(fields: list[dict], title: str = "") -> dict:
    title_field = next((field for field in fields if field.get("key") == "title"), {})
    body_field = next((field for field in fields if field.get("key") == "body"), {})
    return {
        "title": title or title_field.get("defaultValue") or "未命名文档",
        "body": body_field.get("defaultValue") or "",
    }


def normalize_content(content: dict | None, fields: list[dict] | None = None) -> dict:
    base = content_from_fields(fields or [], str((content or {}).get("title") or "")) if fields else {"title": "", "body": ""}
    content = content or {}
    return {
        "title": str(content.get("title") or base.get("title") or ""),
        "body": str(content.get("body") or base.get("body") or ""),
    }


FONT_SIZE_MAP = {
    "初号": 42,
    "小初": 36,
    "一号": 26,
    "小一": 24,
    "二号": 22,
    "小二": 18,
    "三号": 16,
    "小三": 15,
    "四号": 14,
    "小四": 12,
    "五号": 10.5,
    "小五": 9,
}


def _apply_run_font(run, font_name: str | None, font_size: str | None) -> None:
    if font_name:
        run.font.name = font_name
        run._element.rPr.rFonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", font_name)
    if font_size:
        size = FONT_SIZE_MAP.get(font_size)
        if size is None:
            try:
                size = float(font_size)
            except ValueError:
                size = None
        if size:
            run.font.size = Pt(size)


def _apply_paragraph_format(paragraph, field_key: str, format_config: dict | None) -> None:
    format_config = format_config or {}
    font_name = format_config.get("titleFont") if field_key == "title" else format_config.get("bodyFont")
    font_size = (
        format_config.get("titleFontSize")
        if field_key == "title"
        else format_config.get("bodyFontSize")
    ) or format_config.get("fontSize")
    for run in paragraph.runs:
        _apply_run_font(run, font_name, font_size)
    line_spacing = format_config.get("lineSpacing")
    if line_spacing:
        try:
            paragraph.paragraph_format.line_spacing = float(line_spacing)
        except ValueError:
            pass


def _delete_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)
        paragraph._p = paragraph._element = None


def _replace_paragraph_with_lines(paragraph, text: str, field_key: str, format_config: dict | None) -> None:
    lines = str(text or "").replace("\r\n", "\n").replace("\r", "\n").split("\n") or [""]
    paragraph.text = lines[0]
    _apply_paragraph_format(paragraph, field_key, format_config)
    current = paragraph
    for line in lines[1:]:
        current = current.insert_paragraph_before(line)
        _apply_paragraph_format(current, field_key, format_config)


def _replace_xml_block_in_group(paragraphs, field_key: str, text: str, format_config: dict | None) -> bool:
    start_pattern = re.compile(rf"<{field_key}\b[^>]*>", re.IGNORECASE)
    end_pattern = re.compile(rf"</{field_key}>", re.IGNORECASE)
    start_index = None
    end_index = None
    for index, paragraph in enumerate(paragraphs):
        if start_index is None and start_pattern.search(paragraph.text):
            start_index = index
        if start_index is not None and end_pattern.search(paragraph.text):
            end_index = index
            break
    if start_index is None:
        return False
    if end_index is None:
        end_index = start_index

    first = paragraphs[start_index]
    before = start_pattern.split(first.text, maxsplit=1)[0]
    last = paragraphs[end_index]
    after_parts = end_pattern.split(last.text, maxsplit=1)
    after = after_parts[1] if len(after_parts) > 1 else ""
    replacement = f"{before}{text}{after}".strip("\n")
    _replace_paragraph_with_lines(first, replacement, field_key, format_config)

    for paragraph in list(paragraphs[start_index + 1 : end_index + 1]):
        _delete_paragraph(paragraph)
    return True


def _replace_placeholder(document: Document, field_key: str, text: str, format_config: dict | None) -> bool:
    replaced = False
    for paragraph in _iter_paragraphs(document):
        if f"{{{{{field_key}}}}}" in paragraph.text:
            paragraph.text = paragraph.text.replace(f"{{{{{field_key}}}}}", text or "")
            _apply_paragraph_format(paragraph, field_key, format_config)
            replaced = True
    return replaced


def render_document(template_path: str, content: dict, title: str, format_config: dict | None = None) -> tuple[str, str]:
    """按模板替换 title/body 并生成最终 Word 文件。"""
    ensure_writing_storage()
    document = Document(template_path)
    data = normalize_content(content)
    if title and not data.get("title"):
        data["title"] = title

    for field_key in ("title", "body"):
        text = data.get(field_key, "")
        replaced = False
        for paragraphs in _paragraph_groups(document):
            if _replace_xml_block_in_group(paragraphs, field_key, text, format_config):
                replaced = True
        if not replaced:
            _replace_placeholder(document, field_key, text, format_config)

    output_path = DOCUMENT_STORAGE / f"{new_id('wdocfile')}_{data.get('title') or title or 'document'}.docx"
    document.save(output_path)
    return str(output_path), file_sha256(output_path)


def render_blank_document(content: dict, title: str, format_config: dict | None = None) -> tuple[str, str]:
    """无模板空白文稿导出：创建普通 Word 文件并写入标题/正文。"""
    ensure_writing_storage()
    data = normalize_content(content)
    document = Document()
    if data.get("title") or title:
        title_paragraph = document.add_paragraph(data.get("title") or title)
        _apply_paragraph_format(title_paragraph, "title", format_config)
    for line in (data.get("body") or "").replace("\r\n", "\n").replace("\r", "\n").split("\n"):
        paragraph = document.add_paragraph(line)
        _apply_paragraph_format(paragraph, "body", format_config)
    output_path = DOCUMENT_STORAGE / f"{new_id('wdocfile')}_{data.get('title') or title or 'blank'}.docx"
    document.save(output_path)
    return str(output_path), file_sha256(output_path)
